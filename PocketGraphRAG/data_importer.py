"""
PocketGraphRAG 多数据源导入模块

支持从多种数据源提取文本内容：
- 纯文本 (.txt)
- Markdown (.md)
- PDF (.pdf)
- 网页 URL (http/https)
- 批量文件 / 混合导入
"""

import os
import re
from dataclasses import dataclass
from typing import Any, Dict, List, Optional

from .logging_config import get_logger

logger = get_logger(__name__)


def _table_to_markdown(table: List[List[Optional[str]]]) -> str:
    """将 pdfplumber 表格转为 Markdown 格式（对标 LightRAG RAG-Anything 表格序列化）

    Args:
        table: 二维列表，第一行视为表头

    Returns:
        Markdown 表格字符串，空表返回空字符串
    """
    if not table or len(table) < 2:
        return ""
    # 清理 cell：None → 空，去掉换行
    clean = [[(c or "").replace("\n", " ").strip() for c in row] for row in table]
    header = clean[0]
    body = clean[1:]
    # 列数对齐
    n_cols = max(len(r) for r in clean)
    for r in clean:
        while len(r) < n_cols:
            r.append("")
    md = ["| " + " | ".join(header) + " |"]
    md.append("| " + " | ".join(["---"] * n_cols) + " |")
    for row in body:
        md.append("| " + " | ".join(row) + " |")
    return "\n".join(md)


def _extract_pdf_images_metadata(pdf_path: str) -> int:
    """提取 PDF 内嵌图片数量（对标 LightRAG RAG-Anything 图片解析）

    使用 PyMuPDF (fitz) 快速扫描图片，仅统计数量不实际导出。
    实际图片内容解析需 VLM，由 vlm_extractor 模块负责。

    Returns:
        图片数量，PyMuPDF 不可用时返回 0
    """
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        total = 0
        for page in doc:
            total += len(page.get_images(full=False))
        doc.close()
        return total
    except ImportError:
        logger.debug("PyMuPDF (fitz) 未安装，跳过 PDF 图片提取")
        return 0
    except Exception as e:
        logger.debug(f"PDF 图片提取异常: {e}")
        return 0


@dataclass
class ExtractedDocument:
    """提取的文档内容"""

    source: str  # 来源：文件名或 URL
    source_type: str  # 类型：txt / md / pdf / url
    title: str
    content: str
    metadata: Dict[str, Any] = None


class DataImporter:
    """多数据源导入器"""

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf"}

    def __init__(self):
        self.metadata = {}

    def import_file(
        self, file_path: str, image_mode: str = "ocr"
    ) -> Optional[ExtractedDocument]:
        """从单个文件导入

        Args:
            file_path: 文件路径
            image_mode: 图片处理模式 - 'ocr'（提取文字）或 'kg'（直接抽取知识）

        Returns:
            ExtractedDocument 或 None（失败时）
        """
        if not os.path.exists(file_path):
            print(f"[WARN] 文件不存在: {file_path}")
            return None

        ext = os.path.splitext(file_path)[1].lower()
        filename = os.path.basename(file_path)

        try:
            if ext in (".txt",):
                return self._import_txt(file_path)
            elif ext in (".md", ".markdown"):
                return self._import_markdown(file_path)
            elif ext in (".pdf",):
                return self._import_pdf(file_path)
            elif ext in (".doc", ".docx"):
                return self._import_word(file_path)
            elif ext in (
                ".jpg",
                ".jpeg",
                ".png",
                ".gif",
                ".webp",
                ".bmp",
                ".tiff",
                ".tif",
            ):
                return self._import_image(file_path, mode=image_mode)
            else:
                print(f"[WARN] 不支持的文件类型: {ext}")
                return None
        except Exception as e:
            print(f"[ERROR] 导入文件失败 {filename}: {e}")
            return None

    def import_url(
        self, url: str, use_playwright: bool = True
    ) -> Optional[ExtractedDocument]:
        """从网页 URL 导入

        Args:
            url: 网页 URL
            use_playwright: 是否使用 Playwright 渲染动态内容（默认开启）

        Returns:
            ExtractedDocument 或 None
        """
        try:
            return self._import_url(url, use_playwright=use_playwright)
        except Exception as e:
            print(f"[ERROR] 导入网页失败 {url}: {e}")
            return None

    def import_sitemap(
        self,
        sitemap_url: str,
        max_urls: int = 50,
        use_playwright: bool = False,
    ) -> List[ExtractedDocument]:
        """从 sitemap.xml 批量导入网页（对标 RAGFlow 多数据源）

        解析标准 sitemap.xml（含 ``<urlset>`` / ``<sitemap index>``），
        逐个抓取 ``<loc>`` 指向的页面并提取文本。

        Args:
            sitemap_url: sitemap.xml 的 URL
            max_urls: 最多抓取的 URL 数量（防止超大站点失控，默认 50）
            use_playwright: 是否用 Playwright 渲染（默认关闭，sitemap 通常静态）

        Returns:
            成功导入的文档列表
        """
        try:
            urls = self._parse_sitemap(sitemap_url, max_urls=max_urls)
        except Exception as e:
            print(f"[ERROR] 解析 sitemap 失败 {sitemap_url}: {e}")
            return []

        if not urls:
            print(f"[WARN] sitemap 中未找到 URL: {sitemap_url}")
            return []

        print(f"[INFO] sitemap 解析到 {len(urls)} 个 URL，开始批量抓取…")
        results: List[ExtractedDocument] = []
        for i, url in enumerate(urls, 1):
            doc = self.import_url(url, use_playwright=use_playwright)
            if doc is not None:
                results.append(doc)
            if i % 10 == 0:
                print(f"[INFO] sitemap 进度: {i}/{len(urls)}")
        print(f"[INFO] sitemap 导入完成: {len(results)}/{len(urls)} 成功")
        return results

    def import_rss(
        self,
        rss_url: str,
        max_items: int = 50,
        use_playwright: bool = False,
    ) -> List[ExtractedDocument]:
        """从 RSS/Atom feed 批量导入文章（对标 RAGFlow 多数据源）

        支持 RSS 2.0（``<item>`` / ``<link>``）和 Atom 1.0（``<entry>`` / ``<link href>``）。
        优先使用 feed 自带的 ``<content:encoded>`` / ``<description>`` / ``<summary>``，
        缺失时回退到抓取 ``<link>`` 指向的网页。

        Args:
            rss_url: RSS/Atom feed 的 URL
            max_items: 最多导入的条目数（默认 50）
            use_playwright: 是否用 Playwright 渲染（默认关闭）

        Returns:
            成功导入的文档列表
        """
        try:
            items = self._parse_rss(rss_url, max_items=max_items)
        except Exception as e:
            print(f"[ERROR] 解析 RSS 失败 {rss_url}: {e}")
            return []

        if not items:
            print(f"[WARN] RSS 中未找到条目: {rss_url}")
            return []

        print(f"[INFO] RSS 解析到 {len(items)} 个条目，开始导入…")
        results: List[ExtractedDocument] = []
        for item in items:
            # 优先用 feed 自带内容，避免逐页抓取
            if item.get("content"):
                results.append(
                    ExtractedDocument(
                        source=item.get("link", item.get("title", "rss-item")),
                        source_type="rss",
                        title=item.get("title", "Untitled"),
                        content=item["content"],
                        metadata={
                            "link": item.get("link", ""),
                            "published": item.get("published", ""),
                        },
                    )
                )
            elif item.get("link"):
                # 内容缺失，回退到抓取链接
                doc = self.import_url(item["link"], use_playwright=use_playwright)
                if doc is not None:
                    # 用 RSS 条目的标题覆盖（更准确）
                    if item.get("title"):
                        doc.title = item["title"]
                    doc.source_type = "rss"
                    doc.metadata = doc.metadata or {}
                    doc.metadata["published"] = item.get("published", "")
                    results.append(doc)
        print(f"[INFO] RSS 导入完成: {len(results)}/{len(items)} 成功")
        return results

    def _parse_sitemap(
        self, sitemap_url: str, max_urls: int = 50
    ) -> List[str]:
        """解析 sitemap.xml，返回 URL 列表

        支持：
        - ``<urlset>`` (标准 sitemap)
        - ``<sitemapindex>`` (sitemap 索引，递归解析子 sitemap)
        """
        import requests
        from xml.etree import ElementTree as ET

        resp = requests.get(
            sitemap_url,
            headers={"User-Agent": "PocketGraphRAG-SitemapBot/1.0"},
            timeout=30,
        )
        resp.raise_for_status()
        root = ET.fromstring(resp.content)

        # 处理 XML namespace（sitemap.xml 常带 xmlns）
        ns = ""
        if root.tag.startswith("{"):
            ns = root.tag.split("}")[0] + "}"

        urls: List[str] = []

        # sitemapindex：递归解析子 sitemap
        sitemap_tags = root.findall(f"{ns}sitemap")
        if sitemap_tags:
            for sm in sitemap_tags:
                loc = sm.find(f"{ns}loc")
                if loc is not None and loc.text and len(urls) < max_urls:
                    sub_urls = self._parse_sitemap(
                        loc.text.strip(), max_urls=max_urls - len(urls)
                    )
                    urls.extend(sub_urls)
                    if len(urls) >= max_urls:
                        break
            return urls[:max_urls]

        # 标准 urlset
        for url_tag in root.findall(f"{ns}url"):
            loc = url_tag.find(f"{ns}loc")
            if loc is not None and loc.text:
                urls.append(loc.text.strip())
                if len(urls) >= max_urls:
                    break

        return urls[:max_urls]

    def _parse_rss(
        self, rss_url: str, max_items: int = 50
    ) -> List[dict]:
        """解析 RSS/Atom feed，返回条目列表

        每个条目：{"title", "link", "content", "published"}
        支持 RSS 2.0 和 Atom 1.0。
        """
        import requests
        from xml.etree import ElementTree as ET

        resp = requests.get(
            rss_url,
            headers={"User-Agent": "PocketGraphRAG-RSSBot/1.0"},
            timeout=30,
        )
        resp.raise_for_status()
        root = ET.fromstring(resp.content)

        # 处理 namespace
        ns = ""
        if root.tag.startswith("{"):
            ns = root.tag.split("}")[0] + "}"

        items: List[dict] = []

        # RSS 2.0: <item> under <channel>
        if root.tag.endswith("rss") or ns:
            channel = root.find(f"{ns}channel")
            item_parent = channel if channel is not None else root
            for item in item_parent.findall(f"{ns}item"):
                title_el = item.find(f"{ns}title")
                link_el = item.find(f"{ns}link")
                desc_el = item.find(f"{ns}description")
                pub_el = item.find(f"{ns}pubDate")
                # content:encoded (RSS 2.0 with content module)
                content_el = item.find(
                    "{http://purl.org/rss/1.0/modules/content/}encoded"
                )

                content = ""
                if content_el is not None and content_el.text:
                    content = self._html_to_text(content_el.text)
                elif desc_el is not None and desc_el.text:
                    content = self._html_to_text(desc_el.text)

                items.append(
                    {
                        "title": title_el.text.strip() if title_el is not None and title_el.text else "",
                        "link": link_el.text.strip() if link_el is not None and link_el.text else "",
                        "content": content,
                        "published": pub_el.text.strip() if pub_el is not None and pub_el.text else "",
                    }
                )
                if len(items) >= max_items:
                    return items

        # Atom 1.0: <entry>
        atom_ns = "{http://www.w3.org/2005/Atom}"
        for entry in root.findall(f"{atom_ns}entry"):
            title_el = entry.find(f"{atom_ns}title")
            # Atom link: <link href="..." rel="alternate"/>
            link_el = entry.find(f"{atom_ns}link[@rel='alternate']")
            if link_el is None:
                link_el = entry.find(f"{atom_ns}link")
            link = link_el.get("href", "") if link_el is not None else ""
            summary_el = entry.find(f"{atom_ns}summary")
            content_el = entry.find(f"{atom_ns}content")
            pub_el = entry.find(f"{atom_ns}published")
            if pub_el is None:
                pub_el = entry.find(f"{atom_ns}updated")

            content = ""
            if content_el is not None and content_el.text:
                content = self._html_to_text(content_el.text)
            elif summary_el is not None and summary_el.text:
                content = self._html_to_text(summary_el.text)

            items.append(
                {
                    "title": title_el.text.strip() if title_el is not None and title_el.text else "",
                    "link": link,
                    "content": content,
                    "published": pub_el.text.strip() if pub_el is not None and pub_el.text else "",
                }
            )
            if len(items) >= max_items:
                return items

        return items

    @staticmethod
    def _html_to_text(html: str) -> str:
        """简单 HTML → 纯文本转换（避免引入 bs4 依赖时的兜底）"""
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(html, "html.parser")
            return soup.get_text(separator="\n", strip=True)
        except ImportError:
            # 无 bs4 时用正则粗略清理
            import re

            text = re.sub(r"<[^>]+>", " ", html)
            text = re.sub(r"&nbsp;", " ", text)
            text = re.sub(r"&amp;", "&", text)
            text = re.sub(r"&lt;", "<", text)
            text = re.sub(r"&gt;", ">", text)
            text = re.sub(r"\s+", " ", text)
            return text.strip()

    def import_batch(
        self,
        file_paths: List[str] = None,
        urls: List[str] = None,
    ) -> List[ExtractedDocument]:
        """批量导入

        Args:
            file_paths: 文件路径列表
            urls: URL 列表

        Returns:
            成功导入的文档列表
        """
        results = []

        if file_paths:
            for path in file_paths:
                doc = self.import_file(path)
                if doc:
                    results.append(doc)

        if urls:
            for url in urls:
                doc = self.import_url(url)
                if doc:
                    results.append(doc)

        return results

    def import_directory(
        self, dir_path: str, recursive: bool = True
    ) -> List[ExtractedDocument]:
        """导入整个目录下的所有支持的文件

        Args:
            dir_path: 目录路径
            recursive: 是否递归子目录

        Returns:
            成功导入的文档列表
        """
        if not os.path.isdir(dir_path):
            print(f"[WARN] 目录不存在: {dir_path}")
            return []

        file_paths = []
        if recursive:
            for root, _, files in os.walk(dir_path):
                for f in files:
                    ext = os.path.splitext(f)[1].lower()
                    if ext in self.SUPPORTED_EXTENSIONS:
                        file_paths.append(os.path.join(root, f))
        else:
            for f in os.listdir(dir_path):
                ext = os.path.splitext(f)[1].lower()
                if ext in self.SUPPORTED_EXTENSIONS:
                    file_paths.append(os.path.join(dir_path, f))

        return self.import_batch(file_paths=file_paths)

    def _import_txt(self, file_path: str) -> ExtractedDocument:
        """导入纯文本文件"""
        with open(file_path, encoding="utf-8", errors="ignore") as f:
            content = f.read()

        filename = os.path.basename(file_path)
        return ExtractedDocument(
            source=filename,
            source_type="txt",
            title=os.path.splitext(filename)[0],
            content=content.strip(),
            metadata={"file_path": file_path, "size": len(content)},
        )

    def _import_markdown(self, file_path: str) -> ExtractedDocument:
        """导入 Markdown 文件"""
        with open(file_path, encoding="utf-8", errors="ignore") as f:
            content = f.read()

        filename = os.path.basename(file_path)

        # 提取标题（第一个 # 标题）
        title = os.path.splitext(filename)[0]
        match = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
        if match:
            title = match.group(1).strip()

        # 清理 Markdown 标记（保留纯文本内容）
        clean_content = self._markdown_to_text(content)

        return ExtractedDocument(
            source=filename,
            source_type="md",
            title=title,
            content=clean_content.strip(),
            metadata={"file_path": file_path, "size": len(content)},
        )

    def _import_word(self, file_path: str) -> ExtractedDocument:
        """导入 Word 文档（.doc / .docx）"""
        filename = os.path.basename(file_path)
        title = os.path.splitext(filename)[0]

        try:
            from docx import Document

            doc = Document(file_path)

            # 提取正文段落
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())

            # 提取表格内容
            tables_content = []
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(cells))
                if table_rows:
                    tables_content.append("\n".join(table_rows))

            # 合并内容
            content_parts = []
            if paragraphs:
                # 第一行作为标题
                if len(paragraphs) > 0 and len(paragraphs[0]) < 100:
                    title = paragraphs[0]
                    paragraphs = paragraphs[1:]
                content_parts.append("\n\n".join(paragraphs))

            if tables_content:
                content_parts.append(
                    "\n\n--- 表格 ---\n\n" + "\n\n".join(tables_content)
                )

            content = "\n\n".join(content_parts)

            return ExtractedDocument(
                source=filename,
                source_type="word",
                title=title,
                content=content.strip(),
                metadata={
                    "file_path": file_path,
                    "size": len(content),
                    "num_paragraphs": len(paragraphs),
                    "num_tables": len(doc.tables),
                },
            )

        except ImportError:
            raise RuntimeError(
                "Word 文档解析需要 python-docx\n请安装: pip install python-docx"
            )

    def _import_pdf(self, file_path: str, enable_ocr: bool = True) -> ExtractedDocument:
        """导入 PDF 文件，支持表格提取 + 图片提取 + 扫描版 OCR

        三层多模态解析（对标 LightRAG 2025.06 RAG-Anything）：
        1. 文本层：pdfplumber/PyPDF2 提取纯文本
        2. 表格层：pdfplumber.extract_tables() 提取表格结构 → Markdown 表格
        3. 图片层：PyMuPDF(fitz) 提取页面内嵌图片 → 可选 VLM 描述
        4. OCR 兜底：扫描版 PDF 用 VLM 做 OCR

        Args:
            file_path: PDF 文件路径
            enable_ocr: 是否启用扫描版 OCR（检测到文字极少时自动触发）
        """
        filename = os.path.basename(file_path)
        title = os.path.splitext(filename)[0]
        content_parts = []
        tables_parts = []
        total_pages = 0
        ocr_used = False
        num_tables = 0
        num_images = 0

        try:
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    # 1. 文本层
                    text = page.extract_text() or ""
                    if text.strip():
                        content_parts.append(f"--- 第 {i + 1} 页 ---\n{text}")

                    # 2. 表格层（对标 LightRAG RAG-Anything 表格解析）
                    try:
                        tables = page.extract_tables() or []
                        for t_idx, table in enumerate(tables):
                            if not table or len(table) < 2:
                                continue
                            md_table = _table_to_markdown(table)
                            if md_table:
                                tables_parts.append(
                                    f"--- 第 {i + 1} 页 表格 {t_idx + 1} ---\n\n{md_table}"
                                )
                                num_tables += 1
                    except Exception as e:
                        logger.debug(f"第 {i + 1} 页表格提取失败: {e}")

        except ImportError:
            try:
                from PyPDF2 import PdfReader

                reader = PdfReader(file_path)
                if reader.metadata and reader.metadata.title:
                    title = reader.metadata.title
                total_pages = len(reader.pages)
                for i, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        content_parts.append(f"--- 第 {i + 1} 页 ---\n{text}")
            except ImportError:
                raise RuntimeError(
                    "PDF 解析需要 pdfplumber 或 PyPDF2\n"
                    "请安装: pip install pdfplumber 或 pip install PyPDF2"
                )

        # 3. 图片层（可选，对标 LightRAG RAG-Anything 图片解析）
        try:
            num_images = _extract_pdf_images_metadata(file_path)
        except Exception as e:
            logger.debug(f"PDF 图片提取失败: {e}")
            num_images = 0

        # 合并文本 + 表格
        all_parts = content_parts + tables_parts
        content = "\n\n".join(all_parts)

        # 检测是否为扫描版 PDF（文字极少）
        is_scanned = False
        if enable_ocr and total_pages > 0:
            avg_chars_per_page = len(content) / max(total_pages, 1)
            # 平均每页少于 50 个字符，认为是扫描版
            if avg_chars_per_page < 50:
                is_scanned = True

        # 4. 扫描版 PDF：用 VLM 做 OCR
        if is_scanned and enable_ocr:
            try:
                from .vlm_extractor import has_vlm, ocr_scanned_pdf

                if has_vlm():
                    print(
                        f"  [PDF] 检测到扫描版 PDF（平均每页 {avg_chars_per_page:.0f} 字符），启动 OCR..."
                    )
                    ocr_text = ocr_scanned_pdf(file_path)
                    if ocr_text and len(ocr_text) > len(content):
                        content = ocr_text
                        ocr_used = True
                        print(f"  [PDF] OCR 完成，提取 {len(ocr_text)} 字符")
            except Exception as e:
                print(f"  [PDF] OCR 失败: {e}")

        return ExtractedDocument(
            source=filename,
            source_type="pdf",
            title=title,
            content=content.strip(),
            metadata={
                "file_path": file_path,
                "size": len(content),
                "num_pages": total_pages,
                "ocr_used": ocr_used,
                "is_scanned": is_scanned,
                "num_tables": num_tables,
                "num_images": num_images,
            },
        )

    def _import_image(self, file_path: str, mode: str = "ocr") -> ExtractedDocument:
        """导入图片文件，支持 OCR 文字提取和知识抽取

        Args:
            file_path: 图片文件路径
            mode: 模式 - 'ocr'（只提取文字）或 'kg'（直接抽取知识三元组）
        """
        from .vlm_extractor import extract_kg_from_image, has_vlm, ocr_image

        filename = os.path.basename(file_path)
        title = os.path.splitext(filename)[0]

        if not has_vlm():
            raise RuntimeError(
                "图片识别需要配置 VLM 后端（DashScope/Ollama VLM 等）\n"
                "请设置 DASHSCOPE_API_KEY 或 OLLAMA_MODEL（VLM模型）"
            )

        if mode == "kg":
            # 直接抽取知识三元组
            triples, summary = extract_kg_from_image(image_path=file_path)
            content = summary or ""
            if triples:
                content += "\n\n## 提取的知识三元组\n\n"
                for t in triples:
                    head = t.get("head", "")
                    rel = t.get("relation", "")
                    tail = t.get("tail", "")
                    conf = t.get("confidence", 0)
                    content += f"- ({head}, {rel}, {tail}) [置信度: {conf}]\n"
        else:
            # OCR 文字提取
            content = ocr_image(image_path=file_path) or ""

        return ExtractedDocument(
            source=filename,
            source_type="image",
            title=title,
            content=content.strip(),
            metadata={
                "file_path": file_path,
                "size": len(content),
                "mode": mode,
            },
        )

    def _import_url(self, url: str, use_playwright: bool = True) -> ExtractedDocument:
        """导入网页 URL，支持 Playwright 渲染动态内容

        Args:
            url: 网页 URL
            use_playwright: 是否尝试用 Playwright 渲染（默认开启，失败自动降级）
        """
        from urllib.parse import urlparse

        title = urlparse(url).netloc
        content = ""
        rendered_with = "requests"

        # 优先尝试 Playwright 渲染动态内容
        if use_playwright:
            try:
                pw_result = self._fetch_with_playwright(url)
                if pw_result and pw_result.get("content", "").strip():
                    content = pw_result["content"]
                    title = pw_result.get("title", title)
                    rendered_with = "playwright"
            except Exception as e:
                logger.debug("Playwright 渲染失败，回退到静态抓取: %s", e)

        # Playwright 不可用或失败，用 requests 静态抓取
        if not content.strip():
            try:
                from urllib.parse import urlparse

                import requests

                parsed = urlparse(url)
                headers = {
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
                    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
                    "Accept-Encoding": "gzip, deflate",
                    "Connection": "keep-alive",
                    "Referer": f"{parsed.scheme}://{parsed.netloc}/",
                    "Upgrade-Insecure-Requests": "1",
                }

                session = requests.Session()
                response = session.get(
                    url, headers=headers, timeout=30, allow_redirects=True
                )
                response.encoding = response.apparent_encoding or "utf-8"
                html = response.text

                try:
                    from bs4 import BeautifulSoup

                    soup = BeautifulSoup(html, "html.parser")

                    if soup.title and soup.title.string:
                        title = soup.title.string.strip()

                    for tag in soup(
                        [
                            "script",
                            "style",
                            "nav",
                            "footer",
                            "header",
                            "aside",
                            "noscript",
                            "iframe",
                            "svg",
                        ]
                    ):
                        tag.decompose()

                    for selector in [
                        ".ad",
                        ".advertisement",
                        ".ads",
                        ".sidebar",
                        "#sidebar",
                        ".cookie-banner",
                        ".modal",
                        ".popup",
                        ".login-wrap",
                        ".comment",
                        ".comments",
                        ".share",
                        ".recommend",
                        ".related",
                        ".hot-news",
                        ".top-news",
                    ]:
                        for el in soup.select(selector):
                            el.decompose()

                    content_selectors = [
                        "article",
                        "main",
                        ".article-content",
                        ".article",
                        "#content",
                        ".content",
                        ".post-content",
                        ".entry-content",
                        ".detail-content",
                        ".news-content",
                        ".text-content",
                        ".detail",
                        ".article-detail",
                        ".post-detail",
                        ".rich_media_content",
                        "#js_content",
                    ]

                    article = None
                    for sel in content_selectors:
                        article = soup.select_one(sel)
                        if article and len(article.get_text(strip=True)) > 100:
                            break

                    if article:
                        content = article.get_text(separator="\n", strip=True)
                    else:
                        body = soup.find("body")
                        if body:
                            content = body.get_text(separator="\n", strip=True)
                        else:
                            content = soup.get_text(separator="\n", strip=True)

                    content = re.sub(r"\n{3,}", "\n\n", content)
                    content = re.sub(r"[ \t]+", " ", content)
                except ImportError:
                    content = re.sub(
                        r"<script[^>]*>.*?</script>", "", html, flags=re.DOTALL
                    )
                    content = re.sub(
                        r"<style[^>]*>.*?</style>", "", content, flags=re.DOTALL
                    )
                    content = re.sub(r"<[^>]+>", "\n", content)
                    content = re.sub(r"\n{3,}", "\n\n", content).strip()

                rendered_with = "requests"
            except Exception as e:
                if not content.strip():
                    raise ValueError(f"网页抓取失败: {e}")

        return ExtractedDocument(
            source=url,
            source_type="url",
            title=title,
            content=content.strip(),
            metadata={
                "url": url,
                "size": len(content),
                "rendered_with": rendered_with,
            },
        )

    def _fetch_with_playwright(
        self,
        url: str,
        wait_selector: str = None,
        scroll_times: int = 3,
        timeout: int = 30000,
    ) -> dict:
        """用 Playwright 渲染动态网页并提取正文

        Args:
            url: 网页 URL
            wait_selector: 等待的 CSS 选择器（可选，默认等 body）
            scroll_times: 滚动次数，触发懒加载
            timeout: 超时时间（毫秒）

        Returns:
            {'title': str, 'content': str} 或 None（失败时）
        """
        try:
            from playwright.sync_api import sync_playwright
        except ImportError:
            return None

        try:
            with sync_playwright() as p:
                browser = p.chromium.launch(headless=True)
                try:
                    context = browser.new_context(
                        user_agent="Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                        viewport={"width": 1920, "height": 1080},
                    )
                    page = context.new_page()
                    page.goto(url, wait_until="domcontentloaded", timeout=timeout)

                    # 等待内容加载
                    try:
                        if wait_selector:
                            page.wait_for_selector(wait_selector, timeout=timeout)
                        else:
                            page.wait_for_load_state(
                                "networkidle", timeout=min(timeout, 15000)
                            )
                    except Exception:
                        # 等不到也没关系，先拿已有内容
                        pass

                    # 滚动页面，触发懒加载
                    for i in range(scroll_times):
                        page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
                        page.wait_for_timeout(800)

                    # 回到顶部
                    page.evaluate("window.scrollTo(0, 0)")

                    # 获取标题
                    title = page.title()

                    # 移除无关元素
                    page.evaluate("""() => {
                        const selectors = ['script', 'style', 'nav', 'footer', 'header', 'aside',
                                          'noscript', 'iframe', 'svg',
                                          '.ad', '.advertisement', '.ads', '.sidebar', '#sidebar',
                                          '.cookie-banner', '.modal', '.popup', '.login-wrap',
                                          '.comment', '.comments', '.share', '.recommend',
                                          '.related', '.hot-news', '.top-news', '.float-layer'];
                        selectors.forEach(sel => {
                            document.querySelectorAll(sel).forEach(el => el.remove());
                        });
                    }""")

                    # 提取正文（按优先级尝试多种选择器）
                    content = page.evaluate("""() => {
                        const selectors = [
                            'article', 'main', '.article-content', '.article',
                            '#content', '.content', '.post-content', '.entry-content',
                            '.detail-content', '.news-content', '.text-content',
                            '.detail', '.article-detail', '.post-detail',
                            '.rich_media_content', '#js_content', '.article-body',
                            '.content-main', '.main-content'
                        ];

                        let bestElement = null;
                        let bestLength = 0;

                        for (const sel of selectors) {
                            const el = document.querySelector(sel);
                            if (el) {
                                const text = el.innerText || '';
                                if (text.trim().length > bestLength) {
                                    bestLength = text.trim().length;
                                    bestElement = el;
                                }
                            }
                        }

                        // 如果找到的正文太短，就用 body
                        if (!bestElement || bestLength < 200) {
                            bestElement = document.body;
                        }

                        return bestElement ? bestElement.innerText : '';
                    }""")

                    # 清理多余空行
                    content = re.sub(r"\n{3,}", "\n\n", content)

                    return {
                        "title": title,
                        "content": content.strip(),
                    }

                finally:
                    browser.close()
        except Exception:
            return None

    @staticmethod
    def _markdown_to_text(md_content: str) -> str:
        """简单的 Markdown 转纯文本"""
        # 移除代码块
        text = re.sub(r"```.*?```", "", md_content, flags=re.DOTALL)
        # 移除行内代码
        text = re.sub(r"`([^`]+)`", r"\1", text)
        # 移除图片语法
        text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
        # 移除链接语法，保留文字
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        # 移除标题标记
        text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
        # 移除加粗/斜体
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"__([^_]+)__", r"\1", text)
        text = re.sub(r"_([^_]+)_", r"\1", text)
        # 移除引用标记
        text = re.sub(r"^>\s+", "", text, flags=re.MULTILINE)
        # 移除列表标记
        text = re.sub(r"^[-*+]\s+", "", text, flags=re.MULTILINE)
        text = re.sub(r"^\d+\.\s+", "", text, flags=re.MULTILINE)
        # 清理多余空行
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


def detect_file_type(file_path: str) -> Optional[str]:
    """检测文件类型"""
    ext = os.path.splitext(file_path)[1].lower()
    type_map = {
        ".txt": "txt",
        ".md": "md",
        ".markdown": "md",
        ".pdf": "pdf",
    }
    return type_map.get(ext)
