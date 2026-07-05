"""多模态解析单元测试（对标 LightRAG 2025.06 RAG-Anything）

验证 PDF 表格提取 + 图片元数据提取 + Word 表格提取。
"""

import os
import tempfile
from typing import List, Optional

import pytest

from PocketGraphRAG.data_importer import (
    DataImporter,
    _table_to_markdown,
    _extract_pdf_images_metadata,
)


# ==========================
# 1. _table_to_markdown
# ==========================


class TestTableToMarkdown:
    """表格转 Markdown 测试"""

    def test_simple_table(self):
        table = [
            ["姓名", "年龄", "城市"],
            ["张三", "25", "北京"],
            ["李四", "30", "上海"],
        ]
        md = _table_to_markdown(table)
        assert "| 姓名 | 年龄 | 城市 |" in md
        assert "| --- | --- | --- |" in md
        assert "| 张三 | 25 | 北京 |" in md
        assert "| 李四 | 30 | 上海 |" in md

    def test_empty_table(self):
        assert _table_to_markdown([]) == ""
        assert _table_to_markdown([["a"]]) == ""  # 只有表头行不算表格

    def test_single_row_not_table(self):
        """只有一行（表头）不算表格"""
        assert _table_to_markdown([["a", "b", "c"]]) == ""

    def test_none_cells(self):
        """None cell 应转为空字符串"""
        table = [
            ["A", "B"],
            [None, "x"],
            ["y", None],
        ]
        md = _table_to_markdown(table)
        # None → 空字符串，竖线间为空
        assert "|  | x |" in md
        assert "| y |  |" in md

    def test_newlines_in_cells(self):
        """cell 内换行应替换为空格"""
        table = [
            ["描述"],
            ["第一行\n第二行"],
        ]
        md = _table_to_markdown(table)
        assert "第一行 第二行" in md

    def test_uneven_columns(self):
        """不等列数应自动补齐"""
        table = [
            ["A", "B", "C"],
            ["1"],  # 只有1列
        ]
        md = _table_to_markdown(table)
        # 应该能正常生成
        assert "| A | B | C |" in md
        assert "| 1 |" in md or "| 1 |   |   |" in md


# ==========================
# 2. _extract_pdf_images_metadata
# ==========================


class TestExtractPdfImages:
    """PDF 图片元数据提取测试"""

    def test_invalid_path_returns_zero(self):
        """无效路径应返回 0 而非抛异常"""
        assert _extract_pdf_images_metadata("nonexistent.pdf") == 0

    def test_text_pdf_no_images(self, tmp_path):
        """纯文本 PDF 应返回 0 张图片"""
        # 用 reportlab 生成简单 PDF，若不可用则跳过
        try:
            from reportlab.pdfgen import canvas
        except ImportError:
            pytest.skip("reportlab not installed")
        pdf_path = str(tmp_path / "text.pdf")
        c = canvas.Canvas(pdf_path)
        c.drawString(100, 750, "Hello World")
        c.save()
        assert _extract_pdf_images_metadata(pdf_path) == 0

    def test_pdf_with_image(self, tmp_path):
        """带图片的 PDF 应返回 >=1"""
        try:
            import fitz
            from reportlab.pdfgen import canvas
            from reportlab.lib.utils import ImageReader
        except ImportError:
            pytest.skip("PyMuPDF or reportlab not installed")

        # 生成简单图片
        img_path = str(tmp_path / "test.png")
        try:
            from PIL import Image
            img = Image.new("RGB", (100, 100), color="red")
            img.save(img_path)
        except ImportError:
            pytest.skip("PIL not installed")

        # 生成带图片的 PDF
        pdf_path = str(tmp_path / "with_img.pdf")
        c = canvas.Canvas(pdf_path)
        c.drawImage(img_path, 100, 500, width=100, height=100)
        c.save()

        assert _extract_pdf_images_metadata(pdf_path) >= 1


# ==========================
# 3. DataImporter PDF 集成测试
# ==========================


class TestDataImporterPDFIntegration:
    """DataImporter PDF 解析集成测试"""

    def test_text_pdf_extraction(self, tmp_path):
        """纯文本 PDF 应能提取文字"""
        try:
            from reportlab.pdfgen import canvas
        except ImportError:
            pytest.skip("reportlab not installed")

        pdf_path = str(tmp_path / "text.pdf")
        c = canvas.Canvas(pdf_path)
        c.drawString(100, 750, "这是一段测试文字")
        c.drawString(100, 700, "GraphRAG 测试")
        c.save()

        importer = DataImporter()
        doc = importer._import_pdf(pdf_path, enable_ocr=False)
        assert doc.source_type == "pdf"
        assert len(doc.content) > 0
        assert doc.metadata["num_pages"] >= 1
        # 纯文本 PDF 不应触发 OCR
        assert doc.metadata["ocr_used"] is False
        assert doc.metadata["is_scanned"] is False
        # 应统计表格和图片数（可能为 0）
        assert "num_tables" in doc.metadata
        assert "num_images" in doc.metadata

    def test_metadata_fields_complete(self, tmp_path):
        """PDF metadata 应包含所有多模态字段"""
        try:
            from reportlab.pdfgen import canvas
        except ImportError:
            pytest.skip("reportlab not installed")

        pdf_path = str(tmp_path / "test.pdf")
        c = canvas.Canvas(pdf_path)
        c.drawString(100, 750, "test")
        c.save()

        importer = DataImporter()
        doc = importer._import_pdf(pdf_path, enable_ocr=False)
        required_fields = [
            "file_path", "size", "num_pages", "ocr_used",
            "is_scanned", "num_tables", "num_images"
        ]
        for field in required_fields:
            assert field in doc.metadata, f"missing field: {field}"


# ==========================
# 4. Word 表格提取（已有功能回归）
# ==========================


class TestWordTableExtraction:
    """Word 表格提取回归测试"""

    def test_word_with_table(self, tmp_path):
        """带表格的 Word 文档应能提取表格内容"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_path = str(tmp_path / "table.docx")
        doc = Document()
        doc.add_paragraph("文档标题")
        # 添加表格
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "姓名"
        table.cell(0, 1).text = "年龄"
        table.cell(1, 0).text = "张三"
        table.cell(1, 1).text = "25"
        table.cell(2, 0).text = "李四"
        table.cell(2, 1).text = "30"
        doc.save(docx_path)

        importer = DataImporter()
        result = importer.import_file(docx_path)
        assert "表格" in result.content or "姓名" in result.content
        assert result.metadata.get("num_tables", 0) >= 1
